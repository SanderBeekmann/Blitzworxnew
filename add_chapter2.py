import docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

doc = docx.Document('BlitzWorx - Fotografie Pakketten.docx')

DARK = RGBColor(0x04, 0x07, 0x11)
CREAM = RGBColor(0xFE, 0xFA, 0xDC)
GREY = RGBColor(0x88, 0x88, 0x88)
TEXT = RGBColor(0x33, 0x33, 0x33)
SUBTLE = RGBColor(0x55, 0x55, 0x55)
DIVIDER_COLOR = RGBColor(0xCA, 0xCA, 0xAA)
FONT = 'Calibri'


def add_run(paragraph, text, size_pt, color=TEXT, bold=False):
    run = paragraph.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    run.font.bold = bold
    return run


def add_empty(doc):
    return doc.add_paragraph()


def add_divider(doc):
    p = doc.add_paragraph()
    add_run(p, '_' * 80, 6, DIVIDER_COLOR)
    return p


def add_title(doc, text, size_pt=28, color=DARK):
    p = doc.add_paragraph()
    add_run(p, text, size_pt, color, bold=True)
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    add_run(p, text, 13, RGBColor(0x66, 0x66, 0x66))
    return p


def add_section_header(doc, text, size_pt=18):
    p = doc.add_paragraph()
    add_run(p, text, size_pt, DARK, bold=True)
    return p


def add_doelgroep(doc, text):
    p = doc.add_paragraph()
    add_run(p, text, 9, GREY)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    for run in p.runs:
        run.clear()
    add_run(p, text, 10, TEXT)
    return p


def add_price_table(doc, name, price):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True

    for cell in table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = parse_xml(
            '<w:shd {} w:fill="040711" w:color="auto"/>'.format(nsdecls("w"))
        )
        tcPr.append(shading)

    cell0 = table.rows[0].cells[0]
    cell0.paragraphs[0].clear()
    add_run(cell0.paragraphs[0], name, 14, CREAM, bold=True)

    cell1 = table.rows[0].cells[1]
    cell1.paragraphs[0].clear()
    add_run(cell1.paragraphs[0], price, 14, CREAM, bold=True)
    cell1.paragraphs[0].alignment = 2

    return table


def add_body_text(doc, text):
    p = doc.add_paragraph()
    add_run(p, text, 10, SUBTLE)
    return p


# --- PAGE BREAK + NEW CHAPTER ---
doc.add_page_break()

add_empty(doc)
add_title(doc, 'FOTOGRAFIE PRIJZEN', 28)
add_subtitle(doc, 'Aangepaste pakketten - april 2026')
add_divider(doc)
add_empty(doc)

add_body_text(
    doc,
    'Onderstaande prijzen zijn herzien op basis van de actuele kostenstructuur. '
    'Alle pakketten inclusief professionele nabewerking en onbeperkt online gebruiksrecht.'
)

# --- STARTER ---
add_empty(doc)
add_price_table(doc, 'STARTER', '\u20ac249')
add_doelgroep(doc, 'MKB, zakelijke dienstverlening, startups')
add_bullet(doc, '1 uur fotografie op locatie')
add_bullet(doc, '1 uur professionele nabewerking')
add_bullet(doc, "15 professionele bewerkte foto's")
add_bullet(doc, 'Portretten, kantoor- en werksfeer')
add_bullet(doc, 'Aflevering bij aflevering website')

# --- HORECA & RETAIL ---
add_empty(doc)
add_price_table(doc, 'HORECA & RETAIL', '\u20ac595')
add_doelgroep(doc, 'Restaurants, cafes, winkels, hotels')
add_bullet(doc, '4 uur fotografie op locatie')
add_bullet(doc, '4 uur professionele nabewerking')
add_bullet(doc, "75 bewerkte foto's")
add_bullet(doc, 'Food, interieur en sfeerbeelden')
add_bullet(doc, 'Social media-formaten inbegrepen')
add_bullet(doc, 'Digitale levering binnen 7 werkdagen')

# --- VASTGOED BASIS ---
add_empty(doc)
add_price_table(doc, 'VASTGOED BASIS', '\u20ac369')
add_doelgroep(doc, 'Makelaars, projectontwikkelaars')
add_bullet(doc, '2 uur fotografie op locatie')
add_bullet(doc, '2 uur professionele nabewerking')
add_bullet(doc, "20 bewerkte foto's")
add_bullet(doc, 'Interieur en exterieur')
add_bullet(doc, 'Levertijd: 3-5 werkdagen')

# --- VASTGOED UITGEBREID ---
add_empty(doc)
add_price_table(doc, 'VASTGOED UITGEBREID', '\u20ac479')
add_doelgroep(doc, 'Makelaars, projectontwikkelaars (grotere panden)')
add_bullet(doc, '3 uur fotografie op locatie')
add_bullet(doc, '3 uur professionele nabewerking')
add_bullet(doc, "30 bewerkte foto's")
add_bullet(doc, 'Interieur en exterieur')
add_bullet(doc, 'Optioneel: drone luchtfoto\u2019s (+\u20ac249)')
add_bullet(doc, 'Levertijd: 3-5 werkdagen')

# --- E-COMMERCE HALVE DAG ---
add_empty(doc)
add_price_table(doc, 'E-COMMERCE HALVE DAG', '\u20ac589')
add_doelgroep(doc, 'Webshops, productbedrijven')
add_bullet(doc, '4 uur fotografie op locatie')
add_bullet(doc, '4 uur professionele nabewerking')
add_bullet(doc, 'Packshots op witte achtergrond')
add_bullet(doc, 'Ca. 16-20 producten')
add_bullet(doc, 'Inclusief social media-formaten')

# --- E-COMMERCE HELE DAG ---
add_empty(doc)
add_price_table(doc, 'E-COMMERCE HELE DAG', '\u20ac1.049')
add_doelgroep(doc, 'Webshops, productbedrijven (groot assortiment)')
add_bullet(doc, '8 uur fotografie op locatie')
add_bullet(doc, '8 uur professionele nabewerking')
add_bullet(doc, 'Packshots op witte achtergrond')
add_bullet(doc, 'Ca. 32-40 producten')
add_bullet(doc, 'Inclusief social media-formaten')
add_bullet(doc, 'Staffelkorting bespreekbaar bij 100+ producten')

# --- VOLLEDIGE DAG SHOOT ---
add_empty(doc)
add_price_table(doc, 'VOLLEDIGE DAG SHOOT', '\u20ac1.049')
add_doelgroep(doc, 'Aannemers, bouwbedrijven, projectontwikkelaars')
add_bullet(doc, '8 uur fotografie op locatie')
add_bullet(doc, '8 uur professionele nabewerking')
add_bullet(doc, "225 bewerkte foto's")

# --- DIVIDER ---
add_empty(doc)
add_divider(doc)

# --- ADD-ONS ---
add_section_header(doc, 'ADD-ONS', 18)
add_doelgroep(doc, 'Toe te voegen aan elk pakket')
add_empty(doc)

add_price_table(doc, 'DRONE FOTOGRAFIE OF VIDEO', '\u20ac249')
add_bullet(doc, '1 uur drone fotografie op locatie')
add_bullet(doc, '1 uur professionele nabewerking')
add_bullet(doc, "10 bewerkte luchtfoto's")
add_bullet(doc, 'Overzicht pand, terrein en omgeving')
add_bullet(doc, 'Upgrade naar drone video: 2 clips van 30 sec i.p.v. fotografie (+\u20ac249)')

add_empty(doc)
add_price_table(doc, 'EXTRA UUR FOTOGRAFIE', '\u20ac119')
add_bullet(doc, '1 extra uur fotografie + 1 uur nabewerking')
add_bullet(doc, 'Toe te voegen aan elk pakket')

add_empty(doc)
add_price_table(doc, 'SPOEDLEVERING', '\u20ac79')
add_bullet(doc, 'Levering binnen 24-48 uur i.p.v. standaard levertijd')

# --- GOED OM TE WETEN ---
add_empty(doc)
add_section_header(doc, 'GOED OM TE WETEN', 14)
add_bullet(doc, 'Alle prijzen zijn exclusief btw')
add_bullet(doc, 'Reiskosten worden apart doorberekend (\u20ac0,23/km buiten regio Purmerend)')
add_bullet(doc, 'Bewerking (editing) is inbegrepen in alle pakketten')
add_bullet(doc, '1 revisieronde inbegrepen, extra revisies op uurtarief')
add_bullet(doc, 'Spoedlevering mogelijk tegen meerprijs')
add_bullet(doc, "Alle foto's worden geleverd in hoge resolutie (JPEG)")
add_bullet(doc, 'Onbeperkt gebruiksrecht voor online en marketing doeleinden')
add_bullet(doc, 'Pakketten zijn combineerbaar met BlitzWorx webdesign diensten')

# --- FOOTER ---
add_empty(doc)
add_empty(doc)
p = doc.add_paragraph()
add_run(p, 'BlitzWorx  |  www.blitzworx.nl  |  info@blitzworx.nl', 10, GREY)

doc.save('BlitzWorx - Fotografie Pakketten v2.docx')
print('Document saved successfully')
